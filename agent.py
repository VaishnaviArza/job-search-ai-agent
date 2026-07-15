import os
from dotenv import load_dotenv
import anthropic
import requests
import json
from bs4 import BeautifulSoup
from datetime import datetime
import re
import hashlib
import uuid
from collections import Counter
from urllib.parse import urlparse
from docx import Document
import io
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()
claude_api_key = os.getenv("ANTHROPIC_API_KEY")

#initialization
client = anthropic.Anthropic(api_key=claude_api_key)

APPLICATIONS_FILE = 'applications.json'
SAVED_APPLICATIONS_DIR = 'saved_applications'
STATUS_OPTIONS = ["Applied", "Shortlisted", "Interview", "Offer", "Rejected"]

CAREER_URL_KEYWORDS = [
    'career', 'careers', 'job', 'jobs',
    'greenhouse.io', 'lever.co', 'myworkdayjobs.com', 'smartrecruiters.com',
    'icims.com', 'ashbyhq.com', 'bamboohr.com', 'workable.com', 'jobvite.com',
]


#Read the word document in resume format
def read_resume(file):
    try:
        doc = Document(io.BytesIO(file.read()))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())
        seen = set()
        clean = []
        for line in text:
            if line not in seen:
                seen.add(line)
                clean.append(line)

        return '\n'.join(clean)
    except Exception as e:
        return f"Error reading resume {e}"


def is_allowed_job_url(url):
    try:
        parsed = urlparse(url)
    except Exception:
        return False
    if parsed.scheme not in ('http', 'https') or not parsed.netloc:
        return False
    domain = parsed.netloc.lower()
    path = parsed.path.lower()
    if 'linkedin.com' in domain and '/jobs' in path:
        return True
    combined = f"{domain}{path}{parsed.query.lower()}"
    return any(keyword in combined for keyword in CAREER_URL_KEYWORDS)


def extract_jd_from_url(url):
    if not is_allowed_job_url(url):
        raise ValueError(
            "Please provide a LinkedIn job posting URL (linkedin.com/jobs/...) "
            "or a recognisable company careers page URL."
        )

    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    current_url = url
    html = None
    try:
        for _ in range(3):
            response = requests.get(current_url, headers=headers, timeout=15)
            response.raise_for_status()
            html = response.text
            match = re.search(
                r'<meta[^>]+http-equiv=["\']refresh["\'][^>]+content=["\'][^"\']*url=([^"\';]+)',
                html, re.IGNORECASE,
            )
            if match:
                next_url = match.group(1).strip()
                if next_url.startswith('/'):
                    parsed = urlparse(current_url)
                    next_url = f"{parsed.scheme}://{parsed.netloc}{next_url}"
                if next_url == current_url:
                    break
                current_url = next_url
                continue
            break
    except requests.RequestException as e:
        raise ValueError(f"Could not fetch the job posting: {e}")

    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup(['script', 'style', 'nav', 'header', 'footer', 'aside', 'noscript']):
        tag.decompose()
    text = soup.get_text(separator='\n')
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    clean_text = '\n'.join(lines)

    if not clean_text:
        raise ValueError("Could not extract any text from that URL. Please paste the description instead.")

    return clean_text[:8000]


def analyze_and_tailor(job_description, resume):
    prompt = f'''
You are an expert at tailoring resumes based on a job description. You have the job description:
{job_description}
You have the user's resume:
{resume}
Perform the following tasks:
1. Identify the key requirements from the job description.
2. Perform a gap analysis between the resume and the job description.
3. Tailor the resume according to the job description.

Gap analysis rules:
1. Identify skills, tools, and experiences the job description requires that are absent or understated in the resume.
2. Present a MAXIMUM of 2 critical gaps.
3. Each gap must be a single line of no more than 6 words, no punctuation-heavy explanations.

Tailored resume rules:
1. Weave job description keywords naturally into existing experience bullets.
2. Do not fabricate experience or change the meaning of any bullet point.
3. Only reword or reframe what already exists to better mirror the job description's language.
4. Preserve the original resume structure and sections exactly.
5. Each work experience entry must have a MAXIMUM of 3 bullet points.
6. Each bullet point must be a MAXIMUM of 2 lines long (roughly 24 words) so the resume fits on a single page.
7. Keep section headers and job lines short so the whole resume comfortably fits on one page.

Output Format:
1. Begin gap analysis with exactly: ## Gap Analysis
2. End gap analysis with exactly: ## End Gap
3. Begin tailored resume with exactly: ## Tailored Resume
4. End tailored resume with exactly: ## End

Inside the gap analysis, output each gap as its own line starting with "- ".

IMPORTANT:
Output the tailored resume using exactly these tags, one per line:
[NAME] full name
[CONTACT] contact details on one line
[SECTION] section header
[JOB] title | company | location | dates
[BULLET] bullet point text without dash or bullet character
[TEXT] plain text
[EDU] degree | institution | dates

Your output should only contain the Gap analysis and the Tailored resume.
'''

    try:
        messages = [{"role": "user", "content": "Tailor my resume for this job"}]
        response = client.messages.create(
            model='claude-sonnet-4-6',
            messages=messages,
            system=prompt,
            max_tokens=2048
        )
        text = response.content[0].text
        gap_analysis = []
        tailored_resume = None

        if '## Gap Analysis' in text and '## End Gap' in text:
            gap_start = text.index('## Gap Analysis') + len("## Gap Analysis")
            gap_end = text.index("## End Gap")
            gap_block = text[gap_start:gap_end].strip()
            for line in gap_block.splitlines():
                line = line.strip().lstrip('-•').strip()
                if line:
                    gap_analysis.append(line)
            gap_analysis = gap_analysis[:2]

        if '## Tailored Resume' in text and '## End' in text:
            resume_start = text.index('## Tailored Resume') + len("## Tailored Resume")
            resume_end = text.index("## End", resume_start)
            tailored_resume = text[resume_start:resume_end].strip()

        if tailored_resume:
            return (gap_analysis, tailored_resume)

        return ([], text)
    except Exception as e:
        return ([], f"Error tailoring resume: {str(e)}")


def generate_cover_letter(job_description, resume, gap_analysis):
    top_gap = gap_analysis[0] if gap_analysis else "no major gap identified"
    prompt = f'''
You are an expert career coach writing a cover letter on behalf of a candidate.

Job Description:
{job_description}

Candidate Resume:
{resume}

Task:
1. Analyze the tone of the job description to judge whether the company culture reads as formal or casual, and match that tone in the letter.
2. Write a complete cover letter consisting of: a brief greeting, exactly 3 body paragraphs, and a brief closing sign-off.
   - Paragraph 1: express interest in the specific role and company, referencing 1-2 specific requirements from the job description.
   - Paragraph 2: highlight the most relevant experience from the resume that matches the job description.
   - Paragraph 3: briefly acknowledge this as an area of growth, then close with enthusiasm: "{top_gap}"
3. Do not fabricate experience.
4. The closing sign-off should be a plain line like "Sincerely," with no name after it.

Output ONLY the greeting, the 3 paragraphs, and the closing, separated by blank lines. No labels, no commentary.
'''
    try:
        messages = [{"role": "user", "content": "Write my cover letter"}]
        response = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=1024,
            messages=messages,
            system=prompt
        )
        text_blocks = [block.text for block in response.content if block.type == "text"]
        result = "\n\n".join(text_blocks).strip()
        return result if result else "Could not generate cover letter."
    except Exception as e:
        return f"Error generating cover letter: {str(e)}"




def generate_linkedin_message(job_description, resume):
    prompt = '''
You are writing a LinkedIn connection request message on behalf of a job applicant.

Rules:
1. 2 to 3 sentences.
2. STRICT maximum of 300 characters total, including spaces.
3. Reference the specific job title/role from the job description.
4. Include one concrete, specific reason the candidate is a strong fit, drawn from the resume.
5. Professional, warm tone. No hashtags, no emojis, no links.
6. Output ONLY the message text -- no quotation marks, no labels, no preamble.
'''
    content = f'''
Job Description:
{job_description}

Candidate Resume:
{resume}
'''
    try:
        messages = [{"role": "user", "content": content}]
        response = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=300,
            messages=messages,
            system=prompt
        )
        text_blocks = [block.text for block in response.content if block.type == "text"]
        message = "\n".join(text_blocks).strip().strip('"')
        if len(message) > 300:
            message = message[:297].rstrip() + "..."
        return message
    except Exception:
        return "Could not generate LinkedIn outreach message."


def extract_job_info(job_description):
    prompt = '''
You are an information extraction assistant. Given a job description, extract the job title,
company name, and location.
Respond with ONLY a single-line JSON object with exactly these keys: "title", "company", "location".
Use "Unknown" for any field that cannot be determined.
'''
    data = {}
    try:
        response = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=300,
            system=prompt,
            messages=[{"role": "user", "content": job_description}]
        )
        text = next((block.text for block in response.content if block.type == "text"), "")
        match = re.search(r'\{.*\}', text, re.DOTALL)
        if match:
            data = json.loads(match.group(0))
    except Exception:
        data = {}

    return {
        "title": data.get("title") or "Unknown",
        "company": data.get("company") or "Unknown",
        "location": data.get("location") or "Unknown",
    }


def mock_interview(user_answer, conversation_history, jd, resume, prior_weak_areas=None):
    prior_weak_areas_block = ""
    if prior_weak_areas:
        areas_text = ", ".join(prior_weak_areas)
        prior_weak_areas_block = f'''
The candidate's most recent mock interview session identified these weak areas: {areas_text}.
Briefly acknowledge these at the start of the interview (e.g. "Last time you struggled with {areas_text} -- let's focus on that today"),
and weight your questions toward these areas in addition to the job description's requirements.
'''

    prompt = f'''
You are a senior hiring manager conducting a job interview.

You have the job description and the candidate's resume.
Conduct a realistic interview following these rules:

1. Start with a brief professional introduction (extract the company name and job roles from the job description)
2. Ask one question at a time.
3. Questions should mix:
   - Behavioural: "Tell me about a time when..."
   - Technical: based on JD requirements
   - Situational: "How would you handle..."
{prior_weak_areas_block}
4. After each answer give brief feedback (2 lines) then ask next question.
   After the feedback and next question, on their own final line, output exactly:
   [SKILLTAGS] tag1, tag2
   where the tags are 1-3 short skill labels relevant to the candidate's answer (e.g. "Communication", "System Design", "Problem Solving").
   Do not include this line on the introductory turn (before any answer has been given).
5. After 5 questions give a comprehensive assessment:
   - Overall performance score out of 10
   - 3 Strongest points
   - 3 Areas to improve
   - Specific advice for this role
6. After delivering the final comprehensive assessment following question 5, end your response with exactly:
## Interview Complete
[SCORE] <integer 1-10>
[WEAKAREAS] area one; area two; area three

Job Description:
{jd}

Candidate Resume:
{resume}

Keep a professional but encouraging tone throughout.
'''
    if user_answer:
        conversation_history.append({"role": "user", "content": user_answer})
    else:
        conversation_history.append({"role": "user", "content": "Please begin the interview"})

    response = client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=2048,
        messages=conversation_history,
        system=prompt
    )
    assistant_message = response.content[0].text
    conversation_history.append({"role": "assistant", "content": assistant_message})

    return assistant_message, conversation_history


def parse_skill_tags(message):
    lines = message.splitlines()
    tags = []
    clean_lines = list(lines)
    for i in range(len(clean_lines) - 1, -1, -1):
        line = clean_lines[i].strip()
        if not line:
            clean_lines.pop(i)
            continue
        if line.startswith('[SKILLTAGS]'):
            tags_text = line[len('[SKILLTAGS]'):].strip()
            tags = [t.strip() for t in tags_text.split(',') if t.strip()]
            clean_lines.pop(i)
        break
    return '\n'.join(clean_lines).strip(), tags


def parse_final_assessment(message):
    score = None
    weak_areas = []
    lines = message.splitlines()
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('## Interview Complete'):
            continue
        if stripped.startswith('[SCORE]'):
            try:
                score = int(re.search(r'\d+', stripped).group())
            except (AttributeError, ValueError):
                score = None
            continue
        if stripped.startswith('[WEAKAREAS]'):
            areas_text = stripped[len('[WEAKAREAS]'):].strip()
            weak_areas = [a.strip() for a in areas_text.split(';') if a.strip()]
            continue
        clean_lines.append(line)
    return '\n'.join(clean_lines).strip(), score, weak_areas


def generate_study_schedule(gap_analysis, weak_areas, days_remaining, busy_periods=None, other_interviews=None):
    if days_remaining < 1:
        return []

    focus_areas = list(dict.fromkeys((gap_analysis or []) + (weak_areas or [])))
    if not focus_areas:
        return []

    focus_text = "\n".join(f"- {area}" for area in focus_areas)
    today_str = datetime.now().strftime("%Y-%m-%d")

    busy_block = ""
    if busy_periods:
        busy_text = "\n".join(
            f"- {b['date']} {b['start']}-{b['end']}: {b.get('summary', 'Busy')}" for b in busy_periods
        )
        busy_block = f'''
Existing calendar commitments (from the user's real Google Calendar, including study
blocks already scheduled for other job applications) -- do NOT propose a block that
overlaps any of these:
{busy_text}
'''

    other_interviews_block = ""
    if other_interviews:
        other_text = "\n".join(
            f"- {o['date']}: {o.get('title', 'Unknown')} at {o.get('company', 'Unknown')}" for o in other_interviews
        )
        other_interviews_block = f'''
The user has other upcoming interviews too. Avoid loading up study blocks on these
exact dates so they aren't overwhelmed:
{other_text}
'''

    prompt = f'''
You are a career coach building a study schedule for an upcoming job interview.

Today's date is {today_str}. The interview is in {days_remaining} day(s).

Focus areas (skill gaps and weak points to address, in priority order):
{focus_text}
{busy_block}{other_interviews_block}
Propose 3 to 6 study blocks between tomorrow and the interview date, each 1-2 hours long,
prioritized by severity/recency of the focus areas above. Spread them across the available days;
do not schedule more than 2 blocks per day. Use realistic times between 08:00 and 21:00.

For EACH block, also recommend exactly ONE reliable learning resource for that block's topic
(a specific website, course, documentation page, or YouTube video). Prefer official documentation
and reputable platforms (Coursera, DataCamp, Udemy, Pluralsight, YouTube).
Do not fabricate a URL; only include one you are confident is real.

Output ONLY in this exact tagged format, one block per line, nothing else:
[BLOCK] YYYY-MM-DD | HH:MM | HH:MM | Topic name | Resource Title | https://url
'''
    try:
        response = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=1536,
            system=prompt,
            messages=[{"role": "user", "content": "Generate my study schedule"}]
        )
        text = next((block.text for block in response.content if block.type == "text"), "")
        schedule = []
        for line in text.splitlines():
            line = line.strip()
            if not line.startswith('[BLOCK]'):
                continue
            parts = [p.strip() for p in line[len('[BLOCK]'):].split('|')]
            if len(parts) < 4:
                continue
            schedule.append({
                "date": parts[0], "start": parts[1], "end": parts[2], "topic": parts[3],
                "resource_title": parts[4] if len(parts) > 4 else "",
                "resource_url": parts[5] if len(parts) > 5 else "",
            })
        return schedule
    except Exception:
        return []


def set_interview_date(app_id, date_str):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            app['interview_date'] = date_str
            app['study_schedule'] = []
    _write_applications(applications)


def save_interview_session(app_id, score, weak_areas, skill_tags):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            sessions = app.setdefault('interview_sessions', [])
            sessions.append({
                "date": datetime.now().isoformat(),
                "score": score,
                "weak_areas": weak_areas,
                "skill_tags": skill_tags,
            })
    _write_applications(applications)


def get_prior_weak_areas(app_id):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            sessions = app.get('interview_sessions') or []
            if sessions:
                return sessions[-1].get('weak_areas', [])
    return []


def get_performance_summary(app_id):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            sessions = app.get('interview_sessions') or []
            scores = [s.get('score') for s in sessions if s.get('score') is not None]
            weak_area_counts = Counter()
            for s in sessions:
                weak_area_counts.update(s.get('weak_areas') or [])
            return {
                "scores": scores,
                "weak_area_counts": dict(weak_area_counts),
                "sessions": sessions,
            }
    return {"scores": [], "weak_area_counts": {}, "sessions": []}


def save_study_schedule(app_id, blocks):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            app['study_schedule'] = blocks
    _write_applications(applications)


def save_calendar_sync(app_id, blocks, event_ids=None):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            app['calendar_synced_blocks'] = blocks
            app['calendar_event_ids'] = event_ids or []
            app['calendar_synced_at'] = datetime.now().isoformat()
    _write_applications(applications)


def get_calendar_sync(app_id):
    for app in load_application():
        if app.get('id') == app_id:
            return app.get('calendar_synced_blocks')
    return None


def get_calendar_event_ids(app_id):
    for app in load_application():
        if app.get('id') == app_id:
            return app.get('calendar_event_ids') or []
    return []


def create_docx(tailored_text):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)

    lines = tailored_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('[NAME]'):
            content = line[6:].strip()
            para = doc.add_paragraph(content)
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(15)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(2)
        elif line.startswith('[CONTACT]'):
            content = line[9:].strip()
            para = doc.add_paragraph(content)
            para.runs[0].font.size = Pt(9.5)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(8)
        elif line.startswith('[SECTION]'):
            content = line[9:].strip()
            para = doc.add_paragraph(content)
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(11.5)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(2)
        elif line.startswith('[JOB]'):
            content = line[5:].strip()
            para = doc.add_paragraph(content)
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(10)
            para.paragraph_format.space_after = Pt(1)
        elif line.startswith('[BULLET]'):
            content = line[8:].strip()
            para = doc.add_paragraph(content, style='List Bullet')
            for run in para.runs:
                run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.0
        elif line.startswith('[TEXT]'):
            content = line[6:].strip()
            para = doc.add_paragraph(content)
            para.paragraph_format.space_after = Pt(2)
        elif line.startswith('[EDU]'):
            content = line[5:].strip()
            para = doc.add_paragraph(content)
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_cover_letter_docx(cover_letter_text):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    paragraphs = [p.strip() for p in cover_letter_text.split('\n\n') if p.strip()]
    for para_text in paragraphs:
        para = doc.add_paragraph(para_text)
        para.paragraph_format.space_after = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def load_application():
    try:
        with open(APPLICATIONS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return []


def _write_applications(applications):
    with open(APPLICATIONS_FILE, 'w', encoding='utf-8') as f:
        json.dump(applications, f, indent=4)


def save_job_application(job_description, gap_analysis, tailored_resume_text,
                          resume_docx_bytes, cover_letter_docx_bytes):
    os.makedirs(SAVED_APPLICATIONS_DIR, exist_ok=True)
    applications = load_application()
    jd_hash = hashlib.sha256(job_description.encode('utf-8')).hexdigest()

    existing = next((app for app in applications if app.get('jd_hash') == jd_hash), None)
    app_id = existing['id'] if existing else uuid.uuid4().hex

    resume_path = os.path.join(SAVED_APPLICATIONS_DIR, f"{app_id}_resume.docx")
    cover_letter_path = os.path.join(SAVED_APPLICATIONS_DIR, f"{app_id}_cover_letter.docx")
    with open(resume_path, 'wb') as f:
        f.write(resume_docx_bytes)
    with open(cover_letter_path, 'wb') as f:
        f.write(cover_letter_docx_bytes)

    job_info = extract_job_info(job_description)

    record = {
        "id": app_id,
        "jd_hash": jd_hash,
        "title": job_info["title"],
        "company": job_info["company"],
        "location": job_info["location"],
        "date_applied": datetime.now().strftime("%Y-%m-%d"),
        "status": existing["status"] if existing else "Applied",
        "jd": job_description,
        "gap_analysis": gap_analysis,
        "resume_path": resume_path,
        "cover_letter_path": cover_letter_path,
        "tailored_resume_text": tailored_resume_text,
    }

    if existing:
        applications = [record if app.get('jd_hash') == jd_hash else app for app in applications]
    else:
        applications.append(record)

    _write_applications(applications)
    return record


def update_application_status(app_id, new_status):
    applications = load_application()
    for app in applications:
        if app.get('id') == app_id:
            app['status'] = new_status
    _write_applications(applications)


EDITABLE_FIELDS = ('title', 'company', 'location')


def update_application_fields(app_id, fields):
    applications = load_application()
    updated = None
    for app in applications:
        if app.get('id') == app_id:
            for key, value in fields.items():
                if key in EDITABLE_FIELDS:
                    app[key] = value
            updated = app
    _write_applications(applications)
    return updated
